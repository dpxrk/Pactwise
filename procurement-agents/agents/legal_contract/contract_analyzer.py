"""Contract Analyzer with AI-powered scoring and clause extraction"""

import re
import json
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
import textstat
import spacy
from collections import defaultdict

from integrations.databases.contract_models import ContractType, RiskLevel


class ContractAnalyzer:
    """
    Advanced contract analysis with:
    - Multi-dimensional scoring
    - Clause extraction and classification
    - Risk identification
    - Readability analysis
    - Missing clause detection
    """
    
    def __init__(self):
        # Load spaCy model for NLP (install with: python -m spacy download en_core_web_sm)
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            # Fallback if spaCy model not installed
            self.nlp = None
        
        # Standard clauses by contract type
        self.standard_clauses = {
            ContractType.MSA: [
                "Definitions", "Scope of Services", "Term and Termination",
                "Payment Terms", "Confidentiality", "Intellectual Property",
                "Warranties", "Indemnification", "Limitation of Liability",
                "Governing Law", "Dispute Resolution", "Force Majeure",
                "Assignment", "Severability", "Entire Agreement"
            ],
            ContractType.SOW: [
                "Project Description", "Deliverables", "Timeline",
                "Acceptance Criteria", "Payment Schedule", "Change Management",
                "Resources", "Dependencies", "Assumptions", "Risks"
            ],
            ContractType.NDA: [
                "Definition of Confidential Information", "Obligations",
                "Exceptions", "Term", "Return of Information",
                "Remedies", "No License", "Governing Law"
            ],
            ContractType.PURCHASE: [
                "Purchase Terms", "Delivery", "Inspection and Acceptance",
                "Title and Risk", "Warranties", "Price and Payment",
                "Cancellation", "Returns", "Indemnification", "Compliance"
            ]
        }
        
        # Risk indicators
        self.risk_indicators = {
            "high_risk": [
                "unlimited liability", "no liability cap", "consequential damages",
                "punitive damages", "personal guarantee", "auto-renewal",
                "exclusive dealing", "non-compete", "ownership transfer",
                "no termination right", "one-sided termination"
            ],
            "medium_risk": [
                "net 90", "net 120", "milestone payment", "acceptance criteria",
                "change order", "price escalation", "audit rights",
                "most favored nation", "exclusivity period"
            ],
            "low_risk": [
                "net 30", "standard warranty", "mutual termination",
                "reasonable efforts", "commercially reasonable"
            ]
        }
        
        # Positive indicators
        self.protection_indicators = [
            "limitation of liability", "liability cap", "mutual indemnification",
            "termination for convenience", "dispute escalation", "mediation first",
            "confidentiality obligations", "data protection", "insurance requirements",
            "performance bond", "warranty disclaimer", "as is", "force majeure"
        ]
    
    async def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from PDF or Word document"""
        
        file_path = Path(file_path)
        
        if file_type.lower() == "pdf":
            return self._extract_text_from_pdf(file_path)
        elif file_type.lower() in ["docx", "doc"]:
            return self._extract_text_from_docx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods for better accuracy"""
        
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            pass
        
        # Fallback to PyPDF2 if pdfplumber fails
        if not text:
            try:
                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
            except:
                pass
        
        return text
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    
    async def analyze_contract(
        self,
        contract_text: str,
        contract_type: ContractType
    ) -> Dict[str, Any]:
        """Perform comprehensive contract analysis"""
        
        # Calculate individual scores
        completeness_score = self._calculate_completeness_score(contract_text, contract_type)
        risk_coverage_score = self._calculate_risk_coverage_score(contract_text)
        legal_compliance_score = self._calculate_legal_compliance_score(contract_text)
        clarity_score = self._calculate_clarity_score(contract_text)
        commercial_protection_score = self._calculate_commercial_protection_score(contract_text)
        
        # Identify key information
        parties = self._extract_parties(contract_text)
        dates = self._extract_dates(contract_text)
        amounts = self._extract_amounts(contract_text)
        
        return {
            "scores": {
                "completeness": completeness_score,
                "risk_coverage": risk_coverage_score,
                "legal_compliance": legal_compliance_score,
                "clarity": clarity_score,
                "commercial_protection": commercial_protection_score
            },
            "key_information": {
                "parties": parties,
                "dates": dates,
                "amounts": amounts
            },
            "word_count": len(contract_text.split()),
            "readability": {
                "flesch_reading_ease": textstat.flesch_reading_ease(contract_text),
                "flesch_kincaid_grade": textstat.flesch_kincaid_grade(contract_text)
            }
        }
    
    def _calculate_completeness_score(
        self,
        contract_text: str,
        contract_type: ContractType
    ) -> float:
        """Calculate how complete the contract is"""
        
        required_clauses = self.standard_clauses.get(
            contract_type,
            self.standard_clauses[ContractType.MSA]  # Default to MSA
        )
        
        found_clauses = 0
        contract_lower = contract_text.lower()
        
        for clause in required_clauses:
            # Check if clause heading or concept is present
            if clause.lower() in contract_lower:
                found_clauses += 1
            elif self._find_similar_clause(clause, contract_text):
                found_clauses += 1
        
        return (found_clauses / len(required_clauses)) * 100 if required_clauses else 50
    
    def _calculate_risk_coverage_score(self, contract_text: str) -> float:
        """Calculate how well risks are covered"""
        
        contract_lower = contract_text.lower()
        score = 50  # Base score
        
        # Check for high-risk indicators (penalize)
        for risk_phrase in self.risk_indicators["high_risk"]:
            if risk_phrase in contract_lower:
                score -= 5
        
        # Check for protection indicators (reward)
        for protection in self.protection_indicators:
            if protection in contract_lower:
                score += 3
        
        # Check for specific risk mitigations
        if "limitation of liability" in contract_lower:
            score += 10
        if "indemnification" in contract_lower:
            score += 5
        if "insurance" in contract_lower:
            score += 5
        if "force majeure" in contract_lower:
            score += 5
        
        return min(100, max(0, score))
    
    def _calculate_legal_compliance_score(self, contract_text: str) -> float:
        """Calculate legal compliance score"""
        
        contract_lower = contract_text.lower()
        score = 60  # Base score
        
        # Check for essential legal elements
        legal_elements = [
            "governing law", "jurisdiction", "entire agreement",
            "severability", "notice", "amendment", "waiver"
        ]
        
        for element in legal_elements:
            if element in contract_lower:
                score += 5
        
        # Check for modern compliance requirements
        if "data protection" in contract_lower or "gdpr" in contract_lower:
            score += 5
        if "anti-corruption" in contract_lower or "fcpa" in contract_lower:
            score += 5
        if "export control" in contract_lower:
            score += 3
        
        return min(100, max(0, score))
    
    def _calculate_clarity_score(self, contract_text: str) -> float:
        """Calculate contract clarity score"""
        
        # Use readability metrics
        flesch_score = textstat.flesch_reading_ease(contract_text)
        
        # Convert Flesch score (0-100, higher is easier) to our scale
        # Professional contracts typically score 20-40
        if flesch_score >= 50:
            clarity = 90  # Very clear
        elif flesch_score >= 40:
            clarity = 80  # Clear
        elif flesch_score >= 30:
            clarity = 70  # Acceptable
        elif flesch_score >= 20:
            clarity = 60  # Complex but normal
        else:
            clarity = 50  # Too complex
        
        # Check for ambiguous terms (penalize)
        ambiguous_terms = [
            "reasonable", "approximately", "substantially",
            "material", "promptly", "from time to time"
        ]
        
        contract_lower = contract_text.lower()
        ambiguity_count = sum(
            contract_lower.count(term) for term in ambiguous_terms
        )
        
        # Reduce score for excessive ambiguity
        if ambiguity_count > 10:
            clarity -= 10
        elif ambiguity_count > 5:
            clarity -= 5
        
        return min(100, max(0, clarity))
    
    def _calculate_commercial_protection_score(self, contract_text: str) -> float:
        """Calculate commercial protection score"""
        
        contract_lower = contract_text.lower()
        score = 50  # Base score
        
        # Check for commercial protections
        commercial_protections = [
            "payment terms", "late payment", "interest",
            "price adjustment", "cancellation", "refund",
            "warranty", "service level", "performance guarantee",
            "liquidated damages", "penalties", "credits"
        ]
        
        for protection in commercial_protections:
            if protection in contract_lower:
                score += 4
        
        # Check for IP protection
        if "intellectual property" in contract_lower:
            score += 8
        if "ownership" in contract_lower and "retain" in contract_lower:
            score += 5
        if "license" in contract_lower:
            score += 3
        
        # Check for exit clauses
        if "termination for convenience" in contract_lower:
            score += 5
        if "30 days notice" in contract_lower or "thirty days notice" in contract_lower:
            score += 3
        
        return min(100, max(0, score))
    
    async def extract_clauses(self, contract_text: str) -> List[Dict[str, Any]]:
        """Extract and classify contract clauses"""
        
        clauses = []
        
        # Split contract into sections (basic approach)
        sections = self._split_into_sections(contract_text)
        
        for section_num, (title, content) in enumerate(sections, 1):
            clause_type = self._classify_clause(title, content)
            risk_level = self._assess_clause_risk(content)
            
            clauses.append({
                "clause_number": str(section_num),
                "title": title,
                "type": clause_type,
                "risk_level": risk_level,
                "word_count": len(content.split()),
                "summary": self._summarize_clause(content)
            })
        
        return clauses
    
    def _split_into_sections(self, contract_text: str) -> List[Tuple[str, str]]:
        """Split contract into sections based on numbering and headings"""
        
        sections = []
        
        # Look for numbered sections (1., 2., etc.) or lettered sections
        pattern = r"^(\d+\.?\s+|\([a-z]\)\s+|[A-Z]+\.\s+)([A-Z][^.]+)"
        
        lines = contract_text.split("\n")
        current_title = "Preamble"
        current_content = []
        
        for line in lines:
            match = re.match(pattern, line)
            if match:
                # Save previous section
                if current_content:
                    sections.append((current_title, "\n".join(current_content)))
                
                # Start new section
                current_title = match.group(2).strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Add last section
        if current_content:
            sections.append((current_title, "\n".join(current_content)))
        
        return sections
    
    def _classify_clause(self, title: str, content: str) -> str:
        """Classify clause type based on title and content"""
        
        title_lower = title.lower()
        content_lower = content.lower()
        
        # Classification based on keywords
        clause_types = {
            "payment": ["payment", "invoice", "fee", "cost", "price"],
            "liability": ["liability", "damages", "indemnif", "limitation"],
            "termination": ["termination", "expir", "cancel", "end"],
            "confidentiality": ["confidential", "proprietary", "non-disclosure"],
            "warranty": ["warrant", "guarantee", "represent"],
            "intellectual_property": ["intellectual property", "ip", "copyright", "patent"],
            "dispute": ["dispute", "arbitration", "mediation", "litigation"],
            "compliance": ["compliance", "regulatory", "law", "legal"],
            "delivery": ["delivery", "shipment", "timeline", "schedule"],
            "force_majeure": ["force majeure", "act of god", "unforeseeable"]
        }
        
        for clause_type, keywords in clause_types.items():
            for keyword in keywords:
                if keyword in title_lower or keyword in content_lower[:200]:
                    return clause_type
        
        return "general"
    
    def _assess_clause_risk(self, content: str) -> str:
        """Assess risk level of a clause"""
        
        content_lower = content.lower()
        
        # Check for high-risk indicators
        high_risk_count = sum(
            1 for risk in self.risk_indicators["high_risk"]
            if risk in content_lower
        )
        
        if high_risk_count >= 2:
            return "high"
        elif high_risk_count == 1:
            return "medium"
        
        # Check for protective terms
        protection_count = sum(
            1 for protection in self.protection_indicators
            if protection in content_lower
        )
        
        if protection_count >= 2:
            return "low"
        
        return "medium"
    
    def _summarize_clause(self, content: str, max_length: int = 100) -> str:
        """Generate brief summary of clause"""
        
        # Take first meaningful sentence
        sentences = content.split(".")
        for sentence in sentences:
            cleaned = sentence.strip()
            if len(cleaned) > 20:  # Skip very short fragments
                if len(cleaned) <= max_length:
                    return cleaned + "."
                else:
                    return cleaned[:max_length-3] + "..."
        
        return content[:max_length-3] + "..." if len(content) > max_length else content
    
    async def identify_missing_clauses(
        self,
        contract_text: str,
        contract_type: ContractType
    ) -> List[str]:
        """Identify important missing clauses"""
        
        required_clauses = self.standard_clauses.get(
            contract_type,
            self.standard_clauses[ContractType.MSA]
        )
        
        missing = []
        contract_lower = contract_text.lower()
        
        for clause in required_clauses:
            clause_lower = clause.lower()
            # Check if clause is present
            if clause_lower not in contract_lower:
                # Double-check with alternative phrasings
                if not self._find_similar_clause(clause, contract_text):
                    missing.append(clause)
        
        # Check for critical clauses regardless of contract type
        critical_clauses = [
            "Limitation of Liability",
            "Indemnification", 
            "Governing Law",
            "Termination"
        ]
        
        for clause in critical_clauses:
            if clause not in missing and clause.lower() not in contract_lower:
                missing.append(f"{clause} (Critical)")
        
        return missing
    
    def _find_similar_clause(self, clause_name: str, contract_text: str) -> bool:
        """Find similar clause using synonyms and variations"""
        
        # Define common synonyms and variations
        synonyms = {
            "Termination": ["termination", "expiration", "end of agreement", "cancellation"],
            "Payment Terms": ["payment", "compensation", "fees", "pricing"],
            "Confidentiality": ["confidential", "non-disclosure", "proprietary", "nda"],
            "Intellectual Property": ["intellectual property", "ip", "ownership", "copyright"],
            "Warranties": ["warrant", "guarantee", "representation", "assurance"],
            "Indemnification": ["indemnif", "hold harmless", "defend", "liability"],
            "Force Majeure": ["force majeure", "act of god", "unforeseeable", "excused performance"],
            "Governing Law": ["governing law", "applicable law", "jurisdiction", "venue"]
        }
        
        clause_synonyms = synonyms.get(clause_name, [clause_name.lower()])
        contract_lower = contract_text.lower()
        
        for synonym in clause_synonyms:
            if synonym in contract_lower:
                return True
        
        return False
    
    def _extract_parties(self, contract_text: str) -> List[str]:
        """Extract party names from contract"""
        
        parties = []
        
        # Look for patterns like "between X and Y"
        between_pattern = r"between\s+([A-Z][^,]+?)\s+(?:and|AND)\s+([A-Z][^,]+?)(?:\s|,|\()"
        matches = re.findall(between_pattern, contract_text, re.IGNORECASE)
        
        for match in matches:
            parties.extend([party.strip() for party in match])
        
        # Look for "Party" definitions
        party_pattern = r"[\"']([^\"']+)[\"']\s*(?:means|refers to|shall mean)"
        matches = re.findall(party_pattern, contract_text)
        parties.extend(matches)
        
        # Remove duplicates and clean
        parties = list(set(party.strip() for party in parties if len(party) > 3))
        
        return parties[:4]  # Return top 4 parties
    
    def _extract_dates(self, contract_text: str) -> List[str]:
        """Extract important dates from contract"""
        
        dates = []
        
        # Common date patterns
        date_patterns = [
            r"\b(\d{1,2}/\d{1,2}/\d{2,4})\b",
            r"\b(\d{1,2}-\d{1,2}-\d{2,4})\b",
            r"\b([A-Z][a-z]+ \d{1,2}, \d{4})\b",
            r"\b(\d{1,2} [A-Z][a-z]+ \d{4})\b"
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, contract_text)
            dates.extend(matches)
        
        # Remove duplicates
        dates = list(set(dates))
        
        return dates[:5]  # Return top 5 dates
    
    def _extract_amounts(self, contract_text: str) -> List[str]:
        """Extract monetary amounts from contract"""
        
        amounts = []
        
        # Currency patterns
        currency_patterns = [
            r"\$[\d,]+(?:\.\d{2})?",
            r"USD\s*[\d,]+(?:\.\d{2})?",
            r"€[\d,]+(?:\.\d{2})?",
            r"£[\d,]+(?:\.\d{2})?",
            r"[\d,]+(?:\.\d{2})?\s*(?:dollars|euros|pounds)"
        ]
        
        for pattern in currency_patterns:
            matches = re.findall(pattern, contract_text, re.IGNORECASE)
            amounts.extend(matches)
        
        # Remove duplicates and sort by amount
        amounts = list(set(amounts))
        
        return amounts[:5]  # Return top 5 amounts