"""Template Parser for extracting and managing template variables"""

import re
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from docx import Document
import PyPDF2


class TemplateParser:
    """
    Template parsing engine with:
    - Variable extraction from templates
    - Variable validation
    - Structure analysis
    - UI field generation
    """
    
    def __init__(self):
        # Variable patterns to identify
        self.variable_patterns = [
            (r"\${([^}]+)}", "dollar_brace"),        # ${variable}
            (r"\{\{([^}]+)\}\}", "double_brace"),    # {{variable}}
            (r"\[([A-Z_]+)\]", "bracket_upper"),     # [VARIABLE]
            (r"<<([^>]+)>>", "angle_bracket"),       # <<variable>>
            (r"__([^_]+)__", "double_underscore"),   # __variable__
        ]
        
        # Common variable types and their properties
        self.variable_types = {
            "name": {"type": "text", "validation": "alpha_space"},
            "address": {"type": "textarea", "validation": "address"},
            "date": {"type": "date", "validation": "date"},
            "amount": {"type": "number", "validation": "currency"},
            "price": {"type": "number", "validation": "currency"},
            "email": {"type": "email", "validation": "email"},
            "phone": {"type": "tel", "validation": "phone"},
            "percentage": {"type": "number", "validation": "percentage"},
            "terms": {"type": "select", "validation": "options"},
            "description": {"type": "textarea", "validation": "text"},
            "number": {"type": "number", "validation": "numeric"},
            "id": {"type": "text", "validation": "alphanumeric"},
            "url": {"type": "url", "validation": "url"}
        }
        
        # Common variable groups for UI organization
        self.variable_groups = {
            "party_information": [
                "party", "name", "address", "contact", "email", "phone", "registration"
            ],
            "project_details": [
                "project", "description", "scope", "deliverables", "timeline", "milestone"
            ],
            "commercial_terms": [
                "price", "amount", "payment", "terms", "schedule", "currency", "discount"
            ],
            "dates": [
                "date", "effective", "expiry", "deadline", "due", "start", "end"
            ],
            "legal": [
                "jurisdiction", "governing_law", "dispute", "arbitration", "court"
            ]
        }
    
    async def extract_variables(
        self,
        template_path: str,
        file_type: str
    ) -> List[Dict[str, Any]]:
        """Extract variables from template document"""
        
        template_path = Path(template_path)
        
        if file_type.lower() in ["docx", "doc"]:
            text = self._extract_text_from_docx(template_path)
        elif file_type.lower() == "pdf":
            text = self._extract_text_from_pdf(template_path)
        else:
            with open(template_path, "r", encoding="utf-8") as f:
                text = f.read()
        
        # Find all variables
        variables = self._find_variables(text)
        
        # Enhance with type information
        enhanced_variables = self._enhance_variables(variables)
        
        # Sort by appearance order
        enhanced_variables.sort(key=lambda x: x.get("position", 0))
        
        return enhanced_variables
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        
        doc = Document(str(file_path))
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Extract from headers/footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                text += paragraph.text + "\n"
            for paragraph in section.footer.paragraphs:
                text += paragraph.text + "\n"
        
        return text
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        
        text = ""
        
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text()
        
        return text
    
    def _find_variables(self, text: str) -> List[Dict[str, Any]]:
        """Find all variables in text"""
        
        variables = []
        seen_keys = set()
        
        for pattern, pattern_type in self.variable_patterns:
            matches = re.finditer(pattern, text)
            
            for match in matches:
                var_name = match.group(1).strip()
                
                # Clean variable name
                var_key = self._clean_variable_key(var_name)
                
                if var_key and var_key not in seen_keys:
                    seen_keys.add(var_key)
                    
                    variables.append({
                        "key": var_key,
                        "original": var_name,
                        "pattern": pattern_type,
                        "position": match.start(),
                        "context": self._get_context(text, match.start(), match.end())
                    })
        
        return variables
    
    def _clean_variable_key(self, var_name: str) -> str:
        """Clean and standardize variable key"""
        
        # Remove special characters and convert to snake_case
        key = re.sub(r"[^\w\s]", "", var_name)
        key = re.sub(r"\s+", "_", key)
        key = key.lower().strip("_")
        
        return key
    
    def _get_context(self, text: str, start: int, end: int, context_size: int = 50) -> str:
        """Get context around variable for better understanding"""
        
        context_start = max(0, start - context_size)
        context_end = min(len(text), end + context_size)
        
        context = text[context_start:context_end]
        
        # Clean up context
        context = " ".join(context.split())
        
        return context
    
    def _enhance_variables(self, variables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Enhance variables with type and validation information"""
        
        enhanced = []
        
        for var in variables:
            var_key = var["key"]
            
            # Determine variable type
            var_type = self._determine_variable_type(var_key, var.get("context", ""))
            
            # Create enhanced variable definition
            enhanced_var = {
                "key": var_key,
                "label": self._generate_label(var_key),
                "type": var_type["type"],
                "required": self._is_required(var_key, var.get("context", "")),
                "validation": var_type.get("validation"),
                "format": var_type.get("format"),
                "placeholder": self._generate_placeholder(var_key, var_type),
                "help_text": self._generate_help_text(var_key, var_type),
                "section": self._determine_section(var_key),
                "order": var.get("position", 0),
                "original_pattern": var.get("pattern"),
                "context": var.get("context")
            }
            
            # Add type-specific properties
            if var_type["type"] == "select":
                enhanced_var["options"] = self._get_default_options(var_key)
            
            if var_type["type"] == "number":
                enhanced_var["min"] = self._get_min_value(var_key)
                enhanced_var["max"] = self._get_max_value(var_key)
            
            enhanced.append(enhanced_var)
        
        return enhanced
    
    def _determine_variable_type(self, var_key: str, context: str) -> Dict[str, Any]:
        """Determine variable type based on key and context"""
        
        var_lower = var_key.lower()
        
        # Check against known types
        for type_keyword, type_info in self.variable_types.items():
            if type_keyword in var_lower:
                return type_info.copy()
        
        # Context-based detection
        context_lower = context.lower()
        
        if any(word in context_lower for word in ["dollars", "$", "usd", "payment", "cost"]):
            return {"type": "number", "validation": "currency", "format": "currency"}
        
        if any(word in context_lower for word in ["date", "day", "month", "year"]):
            return {"type": "date", "validation": "date"}
        
        if any(word in context_lower for word in ["email", "@"]):
            return {"type": "email", "validation": "email"}
        
        if any(word in context_lower for word in ["select", "choose", "option"]):
            return {"type": "select", "validation": "options"}
        
        # Default to text
        return {"type": "text", "validation": "text"}
    
    def _generate_label(self, var_key: str) -> str:
        """Generate human-readable label from variable key"""
        
        # Convert snake_case to Title Case
        words = var_key.split("_")
        
        # Capitalize each word
        label = " ".join(word.capitalize() for word in words)
        
        # Handle common abbreviations
        replacements = {
            "Id": "ID",
            "Url": "URL",
            "Api": "API",
            "Ip": "IP",
            "Po": "PO"
        }
        
        for old, new in replacements.items():
            label = label.replace(old, new)
        
        return label
    
    def _is_required(self, var_key: str, context: str) -> bool:
        """Determine if variable is required"""
        
        # Common required fields
        required_keywords = [
            "party", "name", "date", "amount", "price", "term"
        ]
        
        var_lower = var_key.lower()
        
        for keyword in required_keywords:
            if keyword in var_lower:
                return True
        
        # Check context for requirement indicators
        context_lower = context.lower()
        if "required" in context_lower or "must" in context_lower:
            return True
        
        return False
    
    def _generate_placeholder(self, var_key: str, var_type: Dict[str, Any]) -> str:
        """Generate placeholder text for input field"""
        
        if var_type["type"] == "email":
            return "email@example.com"
        elif var_type["type"] == "date":
            return "YYYY-MM-DD"
        elif var_type["type"] == "tel":
            return "+1 (555) 123-4567"
        elif var_type.get("format") == "currency":
            return "0.00"
        elif "address" in var_key.lower():
            return "123 Main St, City, State ZIP"
        elif "name" in var_key.lower():
            return "Enter full name"
        else:
            return f"Enter {self._generate_label(var_key).lower()}"
    
    def _generate_help_text(self, var_key: str, var_type: Dict[str, Any]) -> str:
        """Generate help text for input field"""
        
        if var_type.get("format") == "currency":
            return "Enter amount in dollars (e.g., 1000.00)"
        elif var_type["type"] == "date":
            return "Select or enter date in YYYY-MM-DD format"
        elif "percentage" in var_key.lower():
            return "Enter percentage value (0-100)"
        elif "email" in var_key.lower():
            return "Enter valid email address"
        elif "phone" in var_key.lower():
            return "Enter phone number with country code"
        else:
            return ""
    
    def _determine_section(self, var_key: str) -> str:
        """Determine which section the variable belongs to"""
        
        var_lower = var_key.lower()
        
        for section, keywords in self.variable_groups.items():
            for keyword in keywords:
                if keyword in var_lower:
                    return section
        
        return "general"
    
    def _get_default_options(self, var_key: str) -> List[str]:
        """Get default options for select fields"""
        
        var_lower = var_key.lower()
        
        if "payment" in var_lower and "terms" in var_lower:
            return ["Net 30", "Net 60", "Net 90", "Due on Receipt", "50% Upfront"]
        elif "currency" in var_lower:
            return ["USD", "EUR", "GBP", "CAD", "AUD"]
        elif "state" in var_lower or "province" in var_lower:
            return ["CA", "NY", "TX", "FL", "IL", "PA", "OH"]  # Common US states
        elif "country" in var_lower:
            return ["United States", "Canada", "United Kingdom", "Australia"]
        elif "status" in var_lower:
            return ["Active", "Pending", "Completed", "Cancelled"]
        else:
            return ["Option 1", "Option 2", "Option 3"]
    
    def _get_min_value(self, var_key: str) -> Optional[float]:
        """Get minimum value for number fields"""
        
        if "percentage" in var_key.lower():
            return 0
        elif "price" in var_key.lower() or "amount" in var_key.lower():
            return 0
        
        return None
    
    def _get_max_value(self, var_key: str) -> Optional[float]:
        """Get maximum value for number fields"""
        
        if "percentage" in var_key.lower():
            return 100
        
        return None
    
    async def validate_variables(
        self,
        template_variables: List[Dict[str, Any]],
        provided_values: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Validate provided values against template requirements"""
        
        errors = []
        warnings = []
        
        for var_def in template_variables:
            var_key = var_def["key"]
            is_required = var_def.get("required", False)
            
            # Check if required variable is provided
            if is_required and var_key not in provided_values:
                errors.append(f"Required field '{var_def['label']}' is missing")
                continue
            
            if var_key in provided_values:
                value = provided_values[var_key]
                
                # Type validation
                validation_type = var_def.get("validation")
                
                if validation_type == "email":
                    if not self._is_valid_email(value):
                        errors.append(f"Invalid email format for '{var_def['label']}'")
                
                elif validation_type == "date":
                    if not self._is_valid_date(value):
                        errors.append(f"Invalid date format for '{var_def['label']}'")
                
                elif validation_type == "currency":
                    if not self._is_valid_currency(value):
                        errors.append(f"Invalid currency format for '{var_def['label']}'")
                
                elif validation_type == "phone":
                    if not self._is_valid_phone(value):
                        warnings.append(f"Phone number format may be invalid for '{var_def['label']}'")
                
                # Range validation for numbers
                if var_def["type"] == "number":
                    try:
                        num_value = float(value)
                        
                        min_val = var_def.get("min")
                        if min_val is not None and num_value < min_val:
                            errors.append(f"Value for '{var_def['label']}' must be at least {min_val}")
                        
                        max_val = var_def.get("max")
                        if max_val is not None and num_value > max_val:
                            errors.append(f"Value for '{var_def['label']}' must not exceed {max_val}")
                    except (ValueError, TypeError):
                        errors.append(f"Invalid number format for '{var_def['label']}'")
        
        return {
            "is_valid": len(errors) == 0,
            "errors": errors,
            "warnings": warnings
        }
    
    def _is_valid_email(self, email: str) -> bool:
        """Validate email format"""
        pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
        return bool(re.match(pattern, str(email)))
    
    def _is_valid_date(self, date: str) -> bool:
        """Validate date format"""
        patterns = [
            r"^\d{4}-\d{2}-\d{2}$",  # YYYY-MM-DD
            r"^\d{2}/\d{2}/\d{4}$",  # MM/DD/YYYY
            r"^\d{2}-\d{2}-\d{4}$",  # MM-DD-YYYY
        ]
        
        for pattern in patterns:
            if re.match(pattern, str(date)):
                return True
        
        return False
    
    def _is_valid_currency(self, value: Any) -> bool:
        """Validate currency format"""
        try:
            # Remove currency symbols and commas
            cleaned = str(value).replace("$", "").replace(",", "").replace("€", "").replace("£", "")
            float(cleaned)
            return True
        except (ValueError, TypeError):
            return False
    
    def _is_valid_phone(self, phone: str) -> bool:
        """Validate phone format (basic)"""
        # Remove common separators
        cleaned = re.sub(r"[\s\-\(\)\.]+", "", str(phone))
        
        # Check if it contains mostly digits
        return len(cleaned) >= 10 and cleaned[0] != "0" and any(c.isdigit() for c in cleaned)
    
    async def analyze_structure(
        self,
        template_path: str,
        file_type: str
    ) -> Dict[str, Any]:
        """Analyze template document structure"""
        
        template_path = Path(template_path)
        
        structure = {
            "sections": [],
            "total_pages": 0,
            "word_count": 0,
            "has_toc": False,
            "has_signature_block": False,
            "formatting": {}
        }
        
        if file_type.lower() in ["docx", "doc"]:
            doc = Document(str(template_path))
            
            # Count pages (approximate)
            structure["total_pages"] = len(doc.paragraphs) // 30  # Rough estimate
            
            # Analyze sections
            current_section = None
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    if current_section:
                        structure["sections"].append(current_section)
                    
                    current_section = {
                        "title": paragraph.text,
                        "level": int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1,
                        "content_length": 0
                    }
                elif current_section:
                    current_section["content_length"] += len(paragraph.text)
            
            if current_section:
                structure["sections"].append(current_section)
            
            # Check for signature block
            text = " ".join(p.text for p in doc.paragraphs)
            structure["has_signature_block"] = "signature" in text.lower() or "_____" in text
            
            # Word count
            structure["word_count"] = len(text.split())
        
        return structure
    
    def enhance_variables_for_ui(
        self,
        variables: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Enhance variables with additional UI properties"""
        
        enhanced = []
        
        for var in variables:
            enhanced_var = var.copy()
            
            # Add UI-specific properties
            enhanced_var["display_order"] = len(enhanced)
            enhanced_var["column_span"] = 2 if var["type"] == "textarea" else 1
            enhanced_var["input_size"] = "large" if var["type"] == "textarea" else "medium"
            
            # Add conditional display logic
            if "address" in var["key"] and "2" in var["key"]:
                enhanced_var["show_if"] = {"key": "has_multiple_addresses", "value": True}
            
            enhanced.append(enhanced_var)
        
        return enhanced
    
    def group_variables_by_section(
        self,
        variables: List[Dict[str, Any]]
    ) -> Dict[str, List[Dict[str, Any]]]:
        """Group variables by section for organized UI display"""
        
        sections = {}
        
        for var in variables:
            section = var.get("section", "general")
            
            if section not in sections:
                sections[section] = {
                    "title": section.replace("_", " ").title(),
                    "variables": [],
                    "description": self._get_section_description(section)
                }
            
            sections[section]["variables"].append(var)
        
        # Sort sections by importance
        section_order = [
            "party_information",
            "project_details",
            "commercial_terms",
            "dates",
            "legal",
            "general"
        ]
        
        ordered_sections = {}
        for section_key in section_order:
            if section_key in sections:
                ordered_sections[section_key] = sections[section_key]
        
        # Add any remaining sections
        for section_key, section_data in sections.items():
            if section_key not in ordered_sections:
                ordered_sections[section_key] = section_data
        
        return ordered_sections
    
    def _get_section_description(self, section: str) -> str:
        """Get description for section"""
        
        descriptions = {
            "party_information": "Information about the parties involved in the contract",
            "project_details": "Specific details about the project or engagement",
            "commercial_terms": "Financial and payment related terms",
            "dates": "Important dates and deadlines",
            "legal": "Legal and compliance related information",
            "general": "Other contract information"
        }
        
        return descriptions.get(section, "Additional contract details")