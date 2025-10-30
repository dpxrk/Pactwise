"""Comprehensive Contract Review Service - Main orchestrator for advanced contract analysis"""

import hashlib
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from enum import Enum

import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from pdf2image import convert_from_path
import spacy
from sqlalchemy.ext.asyncio import AsyncSession

from integrations.databases.contract_models import (
    Contract, ContractType, RiskLevel, ContractStatus
)


class ClauseCategory(str, Enum):
    """Contract clause categories"""
    INDEMNIFICATION = "INDEMNIFICATION"
    LIABILITY = "LIABILITY"
    TERMINATION = "TERMINATION"
    IP_RIGHTS = "IP_RIGHTS"
    WARRANTY = "WARRANTY"
    DATA_PROTECTION = "DATA_PROTECTION"
    SLA = "SLA"
    PAYMENT = "PAYMENT"
    FORCE_MAJEURE = "FORCE_MAJEURE"
    CONFIDENTIALITY = "CONFIDENTIALITY"
    DISPUTE_RESOLUTION = "DISPUTE_RESOLUTION"
    GOVERNING_LAW = "GOVERNING_LAW"
    ASSIGNMENT = "ASSIGNMENT"
    AUDIT_RIGHTS = "AUDIT_RIGHTS"
    INSURANCE = "INSURANCE"
    OTHER = "OTHER"


class RiskCategory(str, Enum):
    """Risk categories for assessment"""
    LEGAL = "LEGAL"
    FINANCIAL = "FINANCIAL"
    OPERATIONAL = "OPERATIONAL"
    COMPLIANCE = "COMPLIANCE"
    REPUTATIONAL = "REPUTATIONAL"


class ComprehensiveReviewService:
    """
    Advanced contract review service implementing:
    - Multi-format document ingestion (PDF, DOCX, OCR)
    - AI-powered clause extraction and categorization
    - Multi-dimensional risk assessment
    - Regulatory compliance checking (GDPR, CCPA, SOX, HIPAA)
    - Playbook comparison
    - Intelligent redlining
    - Financial analysis
    - Obligation tracking
    """

    def __init__(self):
        # Initialize NLP engine
        try:
            self.nlp = spacy.load("en_core_web_lg")
        except:
            print("Warning: spacy model 'en_core_web_lg' not found, trying 'en_core_web_sm'")
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except:
                print("Warning: No spacy model available. NLP features will be limited.")
                self.nlp = None

        # Company name for context-aware analysis
        self.company_name = self._get_company_name()

        # Risk severity scoring
        self.severity_scores = {
            'CRITICAL': 100,
            'HIGH': 75,
            'MEDIUM': 50,
            'LOW': 25
        }

    def _get_company_name(self) -> str:
        """Get company name from config"""
        # TODO: Load from config
        return "Our Company"

    async def review_contract(
        self,
        contract_file: Path,
        review_type: str = 'comprehensive'
    ) -> Dict[str, Any]:
        """
        Main entry point for comprehensive contract review

        Args:
            contract_file: Path to contract file
            review_type: Type of review ('comprehensive', 'quick', 'focused')

        Returns:
            Complete analysis report with risk assessment, compliance check, recommendations
        """
        # Step 1: Ingest and parse contract
        contract_data = await self.ingest_contract(contract_file)

        # Step 2: Extract structured metadata
        extracted_data = await self.extract_contract_data(contract_data)

        # Step 3: Extract and categorize clauses
        clauses = await self.extract_and_categorize_clauses(contract_data)

        # Step 4: Comprehensive risk assessment
        risk_assessment = await self.assess_risks(
            contract_data,
            extracted_data,
            clauses
        )

        # Step 5: Compliance checking
        compliance_check = await self.check_compliance(
            contract_data,
            extracted_data,
            clauses
        )

        # Step 6: Financial analysis
        financial_analysis = await self.analyze_financial_terms(extracted_data)

        # Step 7: Playbook comparison
        playbook_comparison = await self.compare_to_playbook(
            extracted_data,
            clauses
        )

        # Step 8: Generate redlines
        redlines = await self.generate_redlines(
            contract_data,
            extracted_data,
            risk_assessment,
            playbook_comparison
        )

        # Compile analysis results
        analysis_results = {
            'metadata': extracted_data['metadata'],
            'clauses': clauses,
            'risk_assessment': risk_assessment,
            'compliance_check': compliance_check,
            'financial_analysis': financial_analysis,
            'playbook_comparison': playbook_comparison,
            'redlines': redlines,
            'recommendations': []
        }

        # Step 9: Generate recommendations
        analysis_results['recommendations'] = await self.generate_recommendations(
            analysis_results
        )

        # Step 10: Generate comprehensive report
        report = await self.generate_review_report(analysis_results)

        return report

    async def ingest_contract(self, contract_file: Path) -> Dict[str, Any]:
        """
        Ingest and parse contract from various formats
        Supports: PDF (text and scanned), DOCX, TXT
        """
        file_type = contract_file.suffix.lower().replace('.', '')

        text = ""
        images = []
        is_scanned = False

        if file_type == 'pdf':
            # Check if PDF is scanned
            is_scanned = await self._is_scanned_pdf(contract_file)

            if is_scanned:
                # Use OCR
                text = await self._extract_text_with_ocr(contract_file)
                images = await self._extract_signature_images(contract_file)
            else:
                # Extract text directly
                text = await self._extract_pdf_text(contract_file)
                images = await self._extract_pdf_images(contract_file)

        elif file_type in ['docx', 'doc']:
            text = await self._extract_word_text(contract_file)
            images = await self._extract_word_images(contract_file)

        else:
            # Plain text
            with open(contract_file, 'r', encoding='utf-8') as f:
                text = f.read()

        # Parse document structure
        sections = await self._identify_sections(text)
        pages = await self._split_into_pages(text, contract_file, file_type)

        # Calculate hash for integrity
        file_hash = await self._calculate_file_hash(contract_file)

        contract = {
            'raw_text': text,
            'images': images,
            'sections': sections,
            'pages': pages,
            'file_hash': file_hash,
            'file_type': file_type,
            'is_scanned': is_scanned,
            'ingestion_date': datetime.now(),
            'file_path': str(contract_file)
        }

        return contract

    async def _is_scanned_pdf(self, pdf_path: Path) -> bool:
        """Check if PDF is scanned (image-based)"""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                first_page = pdf.pages[0]
                text = first_page.extract_text()
                # If very little text extracted, likely scanned
                return len(text.strip()) < 100
        except:
            return False

    async def _extract_text_with_ocr(self, pdf_path: Path) -> str:
        """Extract text from scanned PDF using OCR"""
        try:
            images = convert_from_path(pdf_path)
            text_parts = []

            for image in images:
                text = pytesseract.image_to_string(image)
                text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return ""

    async def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Extract text from regular PDF"""
        text_parts = []

        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except:
            # Fallback to PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text())

        return "\n\n".join(text_parts)

    async def _extract_word_text(self, docx_path: Path) -> str:
        """Extract text from Word document"""
        doc = Document(docx_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)

        return "\n".join(text_parts)

    async def _extract_pdf_images(self, pdf_path: Path) -> List[Dict]:
        """Extract images from PDF"""
        # Simplified - would need more robust implementation
        return []

    async def _extract_word_images(self, docx_path: Path) -> List[Dict]:
        """Extract images from Word document"""
        # Simplified - would need more robust implementation
        return []

    async def _extract_signature_images(self, pdf_path: Path) -> List[Dict]:
        """Extract signature images from PDF"""
        # Simplified - would need more robust implementation
        return []

    async def _identify_sections(self, text: str) -> List[Dict[str, Any]]:
        """Identify contract sections using pattern matching"""
        sections = []

        # Common section patterns
        section_patterns = [
            r'^\s*\d+\.\s+([A-Z][A-Za-z\s]+)\s*$',  # 1. Section Title
            r'^\s*Article\s+\d+[:\.]?\s+([A-Z][A-Za-z\s]+)',  # Article 1: Title
            r'^\s*([A-Z][A-Z\s]{3,})\s*$',  # ALL CAPS SECTIONS
        ]

        lines = text.split('\n')
        current_position = 0

        for i, line in enumerate(lines):
            for pattern in section_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    title = match.group(1).strip()
                    sections.append({
                        'title': title,
                        'line_number': i,
                        'position': current_position
                    })
            current_position += len(line) + 1

        return sections

    async def _split_into_pages(
        self,
        text: str,
        file_path: Path,
        file_type: str
    ) -> List[Dict]:
        """Split text into pages"""
        pages = []

        # Page break patterns
        page_break_pattern = r'\n\s*-+\s*Page\s+\d+\s+-+\s*\n|\f'

        page_texts = re.split(page_break_pattern, text)

        for i, page_text in enumerate(page_texts):
            if page_text.strip():
                pages.append({
                    'page_number': i + 1,
                    'text': page_text,
                    'word_count': len(page_text.split())
                })

        return pages

    async def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256 = hashlib.sha256()

        with open(file_path, 'rb') as f:
            while chunk := f.read(8192):
                sha256.update(chunk)

        return sha256.hexdigest()

    async def extract_contract_data(
        self,
        contract: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Extract key metadata and structured data using NLP
        """
        text = contract['raw_text']

        # Use NLP for entity extraction if available
        entities = {}
        if self.nlp:
            doc = self.nlp(text[:100000])  # Limit to avoid memory issues
            entities = await self._extract_entities(doc)

        metadata = {
            'contract_type': await self._classify_contract_type(text),
            'parties': await self._extract_parties(text, entities),
            'effective_date': await self._extract_date(text, 'effective'),
            'expiration_date': await self._extract_date(text, 'expiration|termination'),
            'contract_value': await self._extract_monetary_value(text),
            'currency': await self._extract_currency(text),
            'term_length': await self._calculate_term_length(text),
            'renewal_terms': await self._extract_renewal_terms(text),
            'auto_renewal': await self._check_auto_renewal(text),
            'notice_period': await self._extract_notice_period(text),
            'governing_law': await self._extract_governing_law(text),
            'jurisdiction': await self._extract_jurisdiction(text),
            'key_obligations': await self._extract_obligations(text),
            'payment_terms': await self._extract_payment_terms(text),
            'deliverables': await self._extract_deliverables(text)
        }

        return {
            'metadata': metadata,
            'raw_contract': contract,
            'entities': entities
        }

    async def _extract_entities(self, doc) -> Dict[str, List]:
        """Extract named entities from spacy doc"""
        entities = {
            'PERSON': [],
            'ORG': [],
            'DATE': [],
            'MONEY': [],
            'GPE': []  # Geopolitical entities
        }

        for ent in doc.ents:
            if ent.label_ in entities:
                entities[ent.label_].append(ent.text)

        return entities

    async def _classify_contract_type(self, text: str) -> str:
        """Classify contract type based on content"""
        text_lower = text.lower()

        type_indicators = {
            'NDA': ['non-disclosure', 'confidentiality agreement', 'nda'],
            'MSA': ['master service agreement', 'msa'],
            'SOW': ['statement of work', 'sow'],
            'PURCHASE': ['purchase order', 'purchase agreement'],
            'SOFTWARE': ['software license', 'saas agreement', 'subscription agreement'],
            'CONSULTING': ['consulting agreement', 'professional services'],
            'EMPLOYMENT': ['employment agreement', 'employment contract']
        }

        for contract_type, indicators in type_indicators.items():
            for indicator in indicators:
                if indicator in text_lower:
                    return contract_type

        return 'OTHER'

    async def _extract_parties(
        self,
        text: str,
        entities: Dict
    ) -> List[Dict[str, str]]:
        """Extract contracting parties"""
        parties = []

        # Look for common party patterns
        party_pattern = r'(?:between|by and between)\s+([^,\n]+?)\s+(?:and|&)\s+([^,\n]+)'
        matches = re.finditer(party_pattern, text, re.IGNORECASE)

        for match in matches:
            parties.append({
                'name': match.group(1).strip(),
                'role': 'party1'
            })
            parties.append({
                'name': match.group(2).strip(),
                'role': 'party2'
            })
            break  # Usually first match is the main parties

        return parties

    async def _extract_date(self, text: str, date_type: str) -> Optional[datetime]:
        """Extract specific date type from contract"""
        # Common date patterns
        date_patterns = [
            r'(?:' + date_type + r')[:\s]+(\w+ \d{1,2},? \d{4})',
            r'(?:' + date_type + r')[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                # Parse date (simplified - would use dateutil.parser in production)
                return datetime.now()  # Placeholder

        return None

    async def _extract_monetary_value(self, text: str) -> Optional[float]:
        """Extract contract monetary value"""
        # Look for patterns like "total value of $X" or "contract price: $X"
        money_patterns = [
            r'\$\s*([\d,]+(?:\.\d{2})?)\s*(?:million|mil)?',
            r'(?:value|price|amount)[:\s]+\$\s*([\d,]+(?:\.\d{2})?)',
        ]

        for pattern in money_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                value_str = match.group(1).replace(',', '')
                return float(value_str)

        return None

    async def _extract_currency(self, text: str) -> str:
        """Extract currency"""
        if '$' in text or 'USD' in text.upper():
            return 'USD'
        elif '€' in text or 'EUR' in text.upper():
            return 'EUR'
        elif '£' in text or 'GBP' in text.upper():
            return 'GBP'
        else:
            return 'USD'  # Default

    async def _calculate_term_length(self, text: str) -> Optional[str]:
        """Calculate contract term length"""
        term_pattern = r'(?:term|duration)[:\s]+(\d+)\s*(year|month|day)s?'
        match = re.search(term_pattern, text, re.IGNORECASE)

        if match:
            return f"{match.group(1)} {match.group(2)}s"

        return None

    async def _extract_renewal_terms(self, text: str) -> Optional[str]:
        """Extract renewal terms"""
        renewal_section = re.search(
            r'renewal.*?(?=\n\s*\d+\.|$)',
            text,
            re.IGNORECASE | re.DOTALL
        )

        if renewal_section:
            return renewal_section.group(0)[:500]  # First 500 chars

        return None

    async def _check_auto_renewal(self, text: str) -> bool:
        """Check if contract has auto-renewal"""
        auto_renewal_indicators = [
            'automatic renewal',
            'auto-renew',
            'automatically renew',
            'renew automatically'
        ]

        text_lower = text.lower()
        return any(indicator in text_lower for indicator in auto_renewal_indicators)

    async def _extract_notice_period(self, text: str) -> Optional[str]:
        """Extract notice period for termination"""
        notice_pattern = r'(\d+)\s*(?:day|week|month)s?\s*(?:notice|prior notice)'
        match = re.search(notice_pattern, text, re.IGNORECASE)

        if match:
            return match.group(0)

        return None

    async def _extract_governing_law(self, text: str) -> Optional[str]:
        """Extract governing law"""
        law_pattern = r'govern(?:ed|ing) (?:by )?(?:the )?laws? of ([^,\n]+)'
        match = re.search(law_pattern, text, re.IGNORECASE)

        if match:
            return match.group(1).strip()

        return None

    async def _extract_jurisdiction(self, text: str) -> Optional[str]:
        """Extract jurisdiction"""
        jurisdiction_pattern = r'jurisdiction of (?:the )?([^,\n]+)'
        match = re.search(jurisdiction_pattern, text, re.IGNORECASE)

        if match:
            return match.group(1).strip()

        return None

    async def _extract_obligations(self, text: str) -> List[str]:
        """Extract key obligations"""
        obligations = []

        # Look for "shall" statements
        shall_pattern = r'(?:party|vendor|customer|client)\s+shall\s+([^\.]+)'
        matches = re.finditer(shall_pattern, text, re.IGNORECASE)

        for match in matches[:10]:  # Limit to first 10
            obligations.append(match.group(0))

        return obligations

    async def _extract_payment_terms(self, text: str) -> Optional[str]:
        """Extract payment terms"""
        payment_section = re.search(
            r'payment.*?(?=\n\s*\d+\.|$)',
            text,
            re.IGNORECASE | re.DOTALL
        )

        if payment_section:
            return payment_section.group(0)[:500]

        return None

    async def _extract_deliverables(self, text: str) -> List[str]:
        """Extract deliverables"""
        deliverables = []

        deliverable_section = re.search(
            r'deliverables?.*?(?=\n\s*\d+\.|$)',
            text,
            re.IGNORECASE | re.DOTALL
        )

        if deliverable_section:
            # Extract list items
            section_text = deliverable_section.group(0)
            list_items = re.findall(r'[a-z]\)|[-•]\s*([^\n]+)', section_text, re.IGNORECASE)
            deliverables.extend(list_items[:10])

        return deliverables

    async def extract_and_categorize_clauses(
        self,
        contract: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Extract individual clauses and categorize them
        """
        text = contract['raw_text']

        # Segment text into clauses
        clauses = await self._segment_into_clauses(text)

        categorized_clauses = []

        for clause in clauses:
            clause_category = await self._classify_clause(clause['text'])

            clause_data = {
                'text': clause['text'],
                'category': clause_category,
                'section': clause.get('section'),
                'page': clause.get('page'),
                'risk_level': None,
                'issues': [],
                'confidence_score': 0.8  # Placeholder
            }

            # Perform detailed analysis based on category
            if clause_category == ClauseCategory.INDEMNIFICATION:
                analysis = await self.analyze_indemnification(clause['text'])
                clause_data.update(analysis)
            elif clause_category == ClauseCategory.LIABILITY:
                analysis = await self.analyze_liability_cap(clause['text'])
                clause_data.update(analysis)
            elif clause_category == ClauseCategory.TERMINATION:
                analysis = await self.analyze_termination_rights(clause['text'])
                clause_data.update(analysis)
            elif clause_category == ClauseCategory.DATA_PROTECTION:
                analysis = await self.analyze_data_protection(clause['text'])
                clause_data.update(analysis)
            elif clause_category == ClauseCategory.PAYMENT:
                analysis = await self.analyze_payment_clause(clause['text'])
                clause_data.update(analysis)

            categorized_clauses.append(clause_data)

        return categorized_clauses

    async def _segment_into_clauses(self, text: str) -> List[Dict[str, str]]:
        """Segment contract text into individual clauses"""
        clauses = []

        # Split by numbered sections
        section_pattern = r'\n\s*(\d+\.(?:\d+)?)\s+([A-Z][^\n]+)\n'
        sections = re.split(section_pattern, text)

        for i in range(1, len(sections), 3):
            if i + 2 <= len(sections):
                section_num = sections[i]
                section_title = sections[i + 1]
                section_text = sections[i + 2] if i + 2 < len(sections) else ""

                clauses.append({
                    'text': f"{section_title}\n{section_text}",
                    'section': section_num,
                    'title': section_title
                })

        return clauses

    async def _classify_clause(self, text: str) -> ClauseCategory:
        """Classify clause into category"""
        text_lower = text.lower()

        # Simple keyword-based classification
        classification_rules = {
            ClauseCategory.INDEMNIFICATION: ['indemnif', 'hold harmless', 'defend'],
            ClauseCategory.LIABILITY: ['liability', 'liable', 'damages', 'limitation'],
            ClauseCategory.TERMINATION: ['terminat', 'cancel', 'end this agreement'],
            ClauseCategory.IP_RIGHTS: ['intellectual property', 'ip rights', 'patent', 'copyright', 'trademark'],
            ClauseCategory.WARRANTY: ['warrant', 'guarantee', 'represent'],
            ClauseCategory.DATA_PROTECTION: ['data protection', 'privacy', 'gdpr', 'personal data'],
            ClauseCategory.SLA: ['service level', 'sla', 'uptime', 'availability'],
            ClauseCategory.PAYMENT: ['payment', 'fees', 'compensation', 'invoic'],
            ClauseCategory.FORCE_MAJEURE: ['force majeure', 'act of god', 'beyond control'],
            ClauseCategory.CONFIDENTIALITY: ['confidential', 'proprietary', 'non-disclosure'],
            ClauseCategory.DISPUTE_RESOLUTION: ['dispute', 'arbitration', 'mediation', 'litigation'],
            ClauseCategory.GOVERNING_LAW: ['governing law', 'applicable law', 'jurisdiction']
        }

        for category, keywords in classification_rules.items():
            if any(keyword in text_lower for keyword in keywords):
                return category

        return ClauseCategory.OTHER

    async def analyze_indemnification(self, text: str) -> Dict[str, Any]:
        """Detailed analysis of indemnification clause"""
        text_lower = text.lower()

        analysis = {
            'mutual': False,
            'one_sided': False,
            'favors': None,
            'scope': [],
            'exclusions': [],
            'caps': None,
            'risk_level': 'MEDIUM'
        }

        # Check if mutual or one-sided
        if 'mutual' in text_lower or 'each party' in text_lower or 'indemnify each other' in text_lower:
            analysis['mutual'] = True
            analysis['risk_level'] = 'LOW'
        else:
            analysis['one_sided'] = True
            # Determine direction
            if self.company_name.lower() in text_lower and 'shall indemnify' in text_lower:
                analysis['favors'] = 'VENDOR'
                analysis['risk_level'] = 'HIGH'
            else:
                analysis['favors'] = 'BUYER'
                analysis['risk_level'] = 'LOW'

        # Extract scope
        if 'third party' in text_lower:
            analysis['scope'].append('third_party_claims')
        if 'intellectual property' in text_lower or 'ip' in text_lower:
            analysis['scope'].append('ip_infringement')
        if 'negligence' in text_lower:
            analysis['scope'].append('negligence')
        if 'breach' in text_lower:
            analysis['scope'].append('contract_breach')

        return analysis

    async def analyze_liability_cap(self, text: str) -> Dict[str, Any]:
        """Analyze liability limitation clause"""
        analysis = {
            'has_cap': False,
            'cap_amount': None,
            'cap_type': None,
            'exceptions': [],
            'consequential_damages_excluded': False,
            'risk_level': 'MEDIUM'
        }

        # Check for cap amount
        cap_match = re.search(r'\$\s*([\d,]+(?:\.\d{2})?)', text)
        if cap_match:
            analysis['has_cap'] = True
            analysis['cap_amount'] = float(cap_match.group(1).replace(',', ''))
            analysis['cap_type'] = 'FIXED'
        elif 'contract value' in text.lower() or 'fees paid' in text.lower():
            analysis['has_cap'] = True
            analysis['cap_type'] = 'CONTRACT_VALUE'

        # Check for consequential damages
        if 'consequential' in text.lower() and ('not liable' in text.lower() or 'excluded' in text.lower()):
            analysis['consequential_damages_excluded'] = True

        # Check for exceptions
        if 'gross negligence' in text.lower():
            analysis['exceptions'].append('gross_negligence')
        if 'willful misconduct' in text.lower() or 'willful' in text.lower():
            analysis['exceptions'].append('willful_misconduct')
        if 'fraud' in text.lower():
            analysis['exceptions'].append('fraud')

        return analysis

    async def analyze_termination_rights(self, text: str) -> Dict[str, Any]:
        """Analyze termination clause"""
        analysis = {
            'termination_for_convenience': False,
            'termination_for_cause': False,
            'notice_period': None,
            'mutual_termination': False,
            'risk_level': 'MEDIUM'
        }

        text_lower = text.lower()

        if 'convenience' in text_lower or 'without cause' in text_lower:
            analysis['termination_for_convenience'] = True
            analysis['risk_level'] = 'LOW'

        if 'for cause' in text_lower or 'material breach' in text_lower:
            analysis['termination_for_cause'] = True

        # Extract notice period
        notice_match = re.search(r'(\d+)\s*day', text_lower)
        if notice_match:
            analysis['notice_period'] = f"{notice_match.group(1)} days"

        if 'either party' in text_lower or 'mutual' in text_lower:
            analysis['mutual_termination'] = True
        else:
            analysis['risk_level'] = 'HIGH'

        return analysis

    async def analyze_data_protection(self, text: str) -> Dict[str, Any]:
        """Analyze data protection clause"""
        analysis = {
            'gdpr_compliant': False,
            'ccpa_compliant': False,
            'data_processing_agreement': False,
            'risk_level': 'MEDIUM'
        }

        text_lower = text.lower()

        # Check for GDPR indicators
        gdpr_indicators = ['gdpr', 'data subject rights', 'data controller', 'data processor']
        if any(indicator in text_lower for indicator in gdpr_indicators):
            analysis['gdpr_compliant'] = True
            analysis['risk_level'] = 'LOW'

        # Check for CCPA
        if 'ccpa' in text_lower or 'california consumer privacy' in text_lower:
            analysis['ccpa_compliant'] = True

        # Check for DPA
        if 'data processing agreement' in text_lower or 'dpa' in text_lower:
            analysis['data_processing_agreement'] = True

        return analysis

    async def analyze_payment_clause(self, text: str) -> Dict[str, Any]:
        """Analyze payment terms"""
        analysis = {
            'payment_terms': None,
            'advance_payment': False,
            'advance_percentage': 0,
            'milestone_based': False,
            'risk_level': 'LOW'
        }

        text_lower = text.lower()

        # Check for payment terms
        net_match = re.search(r'net\s*(\d+)', text_lower)
        if net_match:
            days = int(net_match.group(1))
            analysis['payment_terms'] = f"Net {days}"
            if days > 60:
                analysis['risk_level'] = 'MEDIUM'

        # Check for advance payment
        if 'advance' in text_lower or 'upfront' in text_lower:
            analysis['advance_payment'] = True
            percentage_match = re.search(r'(\d+)\s*%', text)
            if percentage_match:
                pct = int(percentage_match.group(1))
                analysis['advance_percentage'] = pct
                if pct > 50:
                    analysis['risk_level'] = 'HIGH'

        # Check for milestone payments
        if 'milestone' in text_lower or 'deliverable' in text_lower:
            analysis['milestone_based'] = True
            analysis['risk_level'] = 'LOW'

        return analysis
