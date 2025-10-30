"""Document Generator for creating contracts from templates"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
import base64

from docx import Document
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate
import pdfkit
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


class DocumentGenerator:
    """
    Document generation engine with:
    - Template variable replacement
    - Multiple format support (Word, PDF)
    - Formatting preservation
    - Dynamic content insertion
    """
    
    def __init__(self):
        # Variable patterns to search for
        self.variable_patterns = [
            r"\${([^}]+)}",           # ${variable}
            r"\{\{([^}]+)\}\}",       # {{variable}}
            r"\[([A-Z_]+)\]",         # [VARIABLE]
            r"<<([^>]+)>>",           # <<variable>>
        ]
        
        # Default styles
        self.styles = getSampleStyleSheet()
        self.custom_styles = self._create_custom_styles()
    
    async def generate_from_template(
        self,
        template_path: str,
        variables: Dict[str, Any],
        output_format: str = "pdf",
        output_dir: Optional[str] = None
    ) -> str:
        """Generate document from template with variable substitution"""
        
        template_path = Path(template_path)
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Determine output path
        if output_dir:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = template_path.parent
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"generated_{timestamp}"
        
        # Process based on template type
        if template_path.suffix.lower() in [".docx", ".doc"]:
            # Generate from Word template
            filled_docx = await self._generate_from_docx(template_path, variables, output_dir, output_name)
            
            if output_format == "pdf":
                # Convert to PDF
                return await self._convert_docx_to_pdf(filled_docx, output_dir, output_name)
            else:
                return str(filled_docx)
        
        elif template_path.suffix.lower() == ".pdf":
            # For PDF templates, we need to extract, replace, and regenerate
            return await self._generate_from_pdf_template(template_path, variables, output_dir, output_name)
        
        else:
            # Plain text template
            return await self._generate_from_text(template_path, variables, output_format, output_dir, output_name)
    
    async def _generate_from_docx(
        self,
        template_path: Path,
        variables: Dict[str, Any],
        output_dir: Path,
        output_name: str
    ) -> Path:
        """Generate document from Word template"""
        
        try:
            # Try using docxtpl for Jinja2-style templates
            doc = DocxTemplate(str(template_path))
            
            # Process variables for docxtpl
            context = self._prepare_context(variables)
            
            # Render the template
            doc.render(context)
            
            # Save the filled document
            output_path = output_dir / f"{output_name}.docx"
            doc.save(str(output_path))
            
            return output_path
            
        except:
            # Fallback to manual replacement
            return await self._generate_from_docx_manual(template_path, variables, output_dir, output_name)
    
    async def _generate_from_docx_manual(
        self,
        template_path: Path,
        variables: Dict[str, Any],
        output_dir: Path,
        output_name: str
    ) -> Path:
        """Manual variable replacement in Word document"""
        
        doc = Document(str(template_path))
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = self._replace_variables(run.text, variables)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = self._replace_variables(run.text, variables)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = self._replace_variables(run.text, variables)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = self._replace_variables(run.text, variables)
        
        # Save the filled document
        output_path = output_dir / f"{output_name}.docx"
        doc.save(str(output_path))
        
        return output_path
    
    def _replace_variables(self, text: str, variables: Dict[str, Any]) -> str:
        """Replace all variable patterns in text"""
        
        if not text:
            return text
        
        result = text
        
        # Try each pattern
        for pattern in self.variable_patterns:
            matches = re.findall(pattern, result)
            for match in matches:
                var_name = match.strip()
                
                # Handle nested variables (e.g., party.name)
                value = self._get_nested_value(variables, var_name)
                
                if value is not None:
                    # Replace the entire pattern
                    for p in self.variable_patterns:
                        result = re.sub(
                            p.replace(r"([^}]+)", re.escape(match)),
                            str(value),
                            result
                        )
        
        return result
    
    def _get_nested_value(self, variables: Dict[str, Any], key: str) -> Any:
        """Get value from nested dictionary using dot notation"""
        
        # Handle direct key
        if key in variables:
            return variables[key]
        
        # Handle nested keys (e.g., "party.name")
        keys = key.split(".")
        value = variables
        
        for k in keys:
            if isinstance(value, dict) and k in value:
                value = value[k]
            else:
                # Also try with underscores
                k_underscore = k.replace("-", "_").replace(" ", "_")
                if isinstance(value, dict) and k_underscore in value:
                    value = value[k_underscore]
                else:
                    return None
        
        return value
    
    def _prepare_context(self, variables: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context for template rendering"""
        
        context = variables.copy()
        
        # Add helper variables
        context.update({
            "today": datetime.now().strftime("%B %d, %Y"),
            "current_year": datetime.now().year,
            "current_date": datetime.now().strftime("%Y-%m-%d")
        })
        
        # Format currency values
        for key, value in context.items():
            if isinstance(value, (int, float)) and "price" in key.lower() or "amount" in key.lower():
                context[f"{key}_formatted"] = f"${value:,.2f}"
        
        # Format dates
        for key, value in context.items():
            if "date" in key.lower() and isinstance(value, str):
                try:
                    date_obj = datetime.fromisoformat(value)
                    context[f"{key}_formatted"] = date_obj.strftime("%B %d, %Y")
                except:
                    pass
        
        return context
    
    async def _convert_docx_to_pdf(
        self,
        docx_path: Path,
        output_dir: Path,
        output_name: str
    ) -> str:
        """Convert Word document to PDF"""
        
        output_path = output_dir / f"{output_name}.pdf"
        
        try:
            # Try using python-docx2pdf (Windows) or unoconv (Linux)
            import subprocess
            
            # Check if LibreOffice is available
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
                capture_output=True
            )
            
            if result.returncode == 0:
                # Rename to desired output name
                generated_pdf = docx_path.with_suffix(".pdf")
                if generated_pdf.exists():
                    generated_pdf.rename(output_path)
                return str(output_path)
        
        except:
            pass
        
        # Fallback: Create PDF from scratch with content
        return await self._create_pdf_from_docx(docx_path, output_path)
    
    async def _create_pdf_from_docx(self, docx_path: Path, output_path: Path) -> str:
        """Create PDF from Word document content"""
        
        doc = Document(str(docx_path))
        
        # Create PDF
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for PDF elements
        story = []
        
        # Add content from Word document
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Determine style based on paragraph style
                if paragraph.style.name.startswith("Heading"):
                    style = self.custom_styles["Heading"]
                else:
                    style = self.custom_styles["Normal"]
                
                p = Paragraph(paragraph.text, style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf.build(story)
        
        return str(output_path)
    
    async def _generate_from_pdf_template(
        self,
        template_path: Path,
        variables: Dict[str, Any],
        output_dir: Path,
        output_name: str
    ) -> str:
        """Generate from PDF template (extract, replace, regenerate)"""
        
        # This is more complex - would need PDF form fields or
        # extract text, replace, and regenerate
        # For now, we'll create a new PDF with the content
        
        output_path = output_dir / f"{output_name}.pdf"
        
        # Extract text from template
        import PyPDF2
        text = ""
        
        with open(template_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text()
        
        # Replace variables
        filled_text = self._replace_variables(text, variables)
        
        # Create new PDF
        return await self._create_pdf_from_text(filled_text, output_path)
    
    async def _generate_from_text(
        self,
        template_path: Path,
        variables: Dict[str, Any],
        output_format: str,
        output_dir: Path,
        output_name: str
    ) -> str:
        """Generate from plain text template"""
        
        # Read template
        with open(template_path, "r", encoding="utf-8") as f:
            template_text = f.read()
        
        # Replace variables
        filled_text = self._replace_variables(template_text, variables)
        
        if output_format == "pdf":
            output_path = output_dir / f"{output_name}.pdf"
            return await self._create_pdf_from_text(filled_text, output_path)
        else:
            # Save as text or HTML
            output_path = output_dir / f"{output_name}.txt"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(filled_text)
            return str(output_path)
    
    async def _create_pdf_from_text(self, text: str, output_path: Path) -> str:
        """Create PDF from plain text"""
        
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        story = []
        
        # Split text into paragraphs
        paragraphs = text.split("\n\n")
        
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a heading (all caps or starts with number)
                if para_text.isupper() or re.match(r"^\d+\.", para_text):
                    style = self.custom_styles["Heading"]
                else:
                    style = self.custom_styles["Normal"]
                
                p = Paragraph(para_text, style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf.build(story)
        
        return str(output_path)
    
    def _create_custom_styles(self) -> Dict[str, ParagraphStyle]:
        """Create custom paragraph styles for PDF"""
        
        custom_styles = {}
        
        # Normal text style
        custom_styles["Normal"] = ParagraphStyle(
            name="CustomNormal",
            parent=self.styles["Normal"],
            fontSize=11,
            leading=14,
            spaceAfter=6,
        )
        
        # Heading style
        custom_styles["Heading"] = ParagraphStyle(
            name="CustomHeading",
            parent=self.styles["Heading1"],
            fontSize=14,
            leading=16,
            spaceAfter=12,
            textColor=RGBColor(0, 0, 0),
        )
        
        # Contract clause style
        custom_styles["Clause"] = ParagraphStyle(
            name="Clause",
            parent=self.styles["Normal"],
            fontSize=11,
            leading=14,
            leftIndent=36,
            spaceAfter=6,
        )
        
        return custom_styles
    
    async def merge_documents(
        self,
        documents: List[str],
        output_path: str,
        output_format: str = "pdf"
    ) -> str:
        """Merge multiple documents into one"""
        
        if output_format == "pdf":
            return await self._merge_pdfs(documents, output_path)
        else:
            return await self._merge_docx(documents, output_path)
    
    async def _merge_pdfs(self, pdf_paths: List[str], output_path: str) -> str:
        """Merge multiple PDFs into one"""
        
        import PyPDF2
        
        pdf_merger = PyPDF2.PdfMerger()
        
        for pdf_path in pdf_paths:
            if os.path.exists(pdf_path):
                pdf_merger.append(pdf_path)
        
        with open(output_path, "wb") as f:
            pdf_merger.write(f)
        
        return output_path
    
    async def _merge_docx(self, docx_paths: List[str], output_path: str) -> str:
        """Merge multiple Word documents into one"""
        
        merged_doc = Document()
        
        for docx_path in docx_paths:
            if os.path.exists(docx_path):
                doc = Document(docx_path)
                
                # Add page break between documents
                if len(merged_doc.paragraphs) > 0:
                    merged_doc.add_page_break()
                
                # Copy content
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
        
        merged_doc.save(output_path)
        
        return output_path